#!/usr/bin/env python3
"""
Generate deterministic binary test fixtures that are impractical to build in JS.

Run once from the backend venv:
    cd backend && source .venv/bin/activate
    python ../frontend/tests/e2e/test-data/generate-binary-fixtures.py

Everything here is small and deterministic (fixed seeds, fixed content) so the
committed bytes never churn. The JS-buildable fixtures (plain multi-page PDFs,
PNGs, text payloads) are generated at runtime by helpers/files.ts instead.
"""
from __future__ import annotations

import io
import os
import struct
import zlib

OUT = os.path.dirname(os.path.abspath(__file__))


def w(name: str, data: bytes) -> None:
    path = os.path.join(OUT, name)
    with open(path, "wb") as f:
        f.write(data)
    print(f"  {name:32s} {len(data):>9,d} bytes")


def make_pdf(pages: int, text: str) -> bytes:
    import fitz

    doc = fitz.open()
    for i in range(pages):
        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 72), f"{text} — page {i + 1}", fontsize=14)
    data = doc.tobytes()
    doc.close()
    return data


def main() -> None:
    import fitz
    from PIL import Image

    print("PDF fixtures:")

    # Encrypted / password-protected PDF (user password = "secret123").
    import pypdf

    base = make_pdf(2, "Locked document")
    reader = pypdf.PdfReader(io.BytesIO(base))
    writer = pypdf.PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    writer.encrypt("secret123", algorithm="AES-256")
    buf = io.BytesIO()
    writer.write(buf)
    w("locked.pdf", buf.getvalue())

    # Image-only / scanned PDF: rasterize a page so there is no text layer.
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    img = Image.new("RGB", (600, 200), "white")
    # Draw crude "text" as black rectangles so OCR has something, but the PDF
    # itself carries no selectable text layer.
    for x in range(20, 560, 40):
        for y in range(80, 120):
            for xx in range(x, x + 20):
                img.putpixel((xx, y), (0, 0, 0))
    bio = io.BytesIO()
    img.save(bio, format="PNG")
    page.insert_image(fitz.Rect(40, 40, 555, 220), stream=bio.getvalue())
    w("scanned.pdf", doc.tobytes())
    doc.close()

    # PDF with a real table (for pdf-to-excel).
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    rows = [["Name", "Qty", "Price"], ["Widget", "3", "9.99"], ["Gadget", "1", "19.50"]]
    y = 100
    for r in rows:
        x = 80
        for cell in r:
            page.insert_text((x, y), cell, fontsize=11)
            x += 150
        # horizontal rule so pdfplumber's line-based detector sees a table
        page.draw_line(fitz.Point(70, y + 6), fitz.Point(530, y + 6))
        y += 40
    for gx in (70, 220, 370, 530):
        page.draw_line(fitz.Point(gx, 88), fitz.Point(gx, y - 34))
    w("table.pdf", doc.tobytes())
    doc.close()

    # Unicode / Bangla text PDF.
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Unicode: café résumé naïve — 日本語 — emoji check", fontsize=12)
    w("unicode.pdf", doc.tobytes())
    doc.close()

    # Very small (1-page minimal) and multi-page.
    w("single.pdf", make_pdf(1, "Single page"))
    w("multipage.pdf", make_pdf(10, "Multi page"))

    # Corrupted PDF: valid %PDF header, garbage body, broken xref.
    w("corrupted.pdf", b"%PDF-1.7\n" + b"\xde\xad\xbe\xef" * 64 + b"\nstartxref\n99999\n%%EOF")

    # Zero-byte PDF.
    w("empty.pdf", b"")

    print("Image fixtures:")

    # Valid JPG / PNG / WebP.
    im = Image.new("RGB", (64, 48), (200, 40, 40))
    for fmt, name in [("JPEG", "photo.jpg"), ("PNG", "photo.png"), ("WEBP", "photo.webp")]:
        bio = io.BytesIO()
        im.save(bio, format=fmt)
        w(name, bio.getvalue())

    # Transparent PNG.
    rgba = Image.new("RGBA", (64, 48), (0, 128, 255, 0))
    for x in range(20, 44):
        for y in range(14, 34):
            rgba.putpixel((x, y), (255, 255, 255, 255))
    bio = io.BytesIO()
    rgba.save(bio, format="PNG")
    w("transparent.png", bio.getvalue())

    # Very wide and very tall.
    bio = io.BytesIO()
    Image.new("RGB", (2000, 40), (10, 120, 10)).save(bio, format="PNG")
    w("wide.png", bio.getvalue())
    bio = io.BytesIO()
    Image.new("RGB", (40, 2000), (10, 10, 120)).save(bio, format="PNG")
    w("tall.png", bio.getvalue())

    # Tiny 1x1.
    bio = io.BytesIO()
    Image.new("RGB", (1, 1), (255, 255, 255)).save(bio, format="PNG")
    w("tiny.png", bio.getvalue())

    # HEIC (real, via pillow-heif).
    try:
        import pillow_heif

        pillow_heif.register_heif_opener()
        bio = io.BytesIO()
        Image.new("RGB", (120, 90), (40, 160, 90)).save(bio, format="HEIF", quality=60)
        w("photo.heic", bio.getvalue())
    except Exception as exc:  # pragma: no cover
        print(f"  (skipped HEIC: {exc})")

    # Corrupted image: PNG magic then garbage.
    w("corrupted.png", b"\x89PNG\r\n\x1a\n" + b"\x00\xff" * 40)

    # Renamed-extension: a real PNG saved with a .jpg name lives as a fixture the
    # test copies; here we ship the truthful bytes and let the test rename.
    # (see helpers/files.ts renamedExtension)

    print("Office fixtures:")

    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Product", "Units", "Revenue"])
    ws.append(["Alpha", 12, 1440])
    ws.append(["Beta", 7, 910])
    bio = io.BytesIO()
    wb.save(bio)
    w("book.xlsx", bio.getvalue())

    import docx

    d = docx.Document()
    d.add_heading("Sample Document", level=1)
    d.add_paragraph("The quick brown fox jumps over the lazy dog.")
    d.add_paragraph("Second paragraph with some more text for conversion.")
    bio = io.BytesIO()
    d.save(bio)
    w("doc.docx", bio.getvalue())

    print("Done.")


if __name__ == "__main__":
    main()
