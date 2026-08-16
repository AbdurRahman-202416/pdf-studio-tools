"""Convert a PDF to a Microsoft Word .docx file.

Used by /api/tools/pdf-to-word. We extract text per page with PyMuPDF
(preserving paragraph order and blank lines) and write paragraphs into
a docx via python-docx. Inline images embedded in the PDF are extracted
and inserted into the output for layout fidelity.

This is a best-effort text-and-image pipeline — not a full layout-mirror.
Complex multi-column or table-heavy PDFs may need light cleanup in Word.
"""
from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image

from app.utils.guards import assert_pdf_page_limit
from app.utils.storage import new_file_id


class PdfToWordError(Exception):
    pass


def _extract_blocks(page: fitz.Page) -> list[dict]:
    """Yield text + image blocks in reading order."""
    blocks = page.get_text("dict")["blocks"]
    return blocks


def pdf_to_docx(file_path: Path) -> tuple[str, Path, dict]:
    assert_pdf_page_limit(file_path)
    output_id = new_file_id()
    out = file_path.parent.parent / "output" / f"{output_id}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pages_processed = 0
    images_inserted = 0
    chars = 0

    try:
        with fitz.open(file_path) as src:
            for page_index in range(src.page_count):
                page = src.load_page(page_index)
                if page_index > 0:
                    doc.add_page_break()
                for block in _extract_blocks(page):
                    btype = block.get("type", 0)
                    if btype == 0:  # text
                        paragraph_text = []
                        for line in block.get("lines", []):
                            line_text = "".join(
                                span.get("text", "") for span in line.get("spans", [])
                            )
                            paragraph_text.append(line_text)
                        text = "\n".join(paragraph_text).strip()
                        if text:
                            doc.add_paragraph(text)
                            chars += len(text)
                    elif btype == 1:  # image
                        try:
                            img_bytes = block.get("image")
                            if not img_bytes:
                                continue
                            with Image.open(io.BytesIO(img_bytes)) as im:
                                if im.mode not in ("RGB", "L"):
                                    im = im.convert("RGB")
                                buf = io.BytesIO()
                                im.save(buf, format="PNG", optimize=True)
                                buf.seek(0)
                                # Cap embedded width at 6 inches for readability.
                                doc.add_picture(buf, width=Inches(min(6.0, im.width / 96)))
                                images_inserted += 1
                        except Exception:
                            continue
                pages_processed += 1

        doc.save(out)
    except PdfToWordError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise PdfToWordError(f"Failed to convert PDF to Word: {exc}") from exc

    return output_id, out, {
        "pages": pages_processed,
        "characters": chars,
        "images": images_inserted,
    }
